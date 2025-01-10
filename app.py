import os
from flask import Flask, render_template, request, send_file, jsonify
import pytesseract
from PIL import Image
import cv2
import numpy as np
from docx import Document

app = Flask(__name__, static_folder='static')

# Set the path for Tesseract OCR executable
pytesseract.pytesseract.tesseract_cmd = r"C:\Users\akash\TRASH_FILES\configuration\tesseract\tesseract.exe"  # Adjust if necessary

# Route for uploading image and processing
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_image():
    if 'image' not in request.files:
        return "No file part", 400

    file = request.files['image']
    if file.filename == '':
        return "No selected file", 400

    # Process the image and extract text and tables
    img = Image.open(file)
    img_cv = np.array(img)

    # Extract text using pytesseract
    text = pytesseract.image_to_string(img, config='--psm 6')

    # Detect and extract table if any using OpenCV
    tables = extract_table(img_cv)

    # If no table detected, return only the text
    if not tables:
        return jsonify({'text': text, 'docx_url': None})
    
    # Generate .docx file with both text and table
    docx_file = create_docx(text, tables)

    # Return the text and generated file URL
    return jsonify({'text': text, 'docx_url': docx_file})

def extract_table(img):
    """ Use OpenCV to detect tables in the image and return the table data. """
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    
    # Apply adaptive thresholding to handle different lighting conditions
    thresh = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                                  cv2.THRESH_BINARY, 11, 2)

    # Detect edges using Canny edge detection
    edges = cv2.Canny(thresh, 50, 150)

    # Find contours which might represent the table grid
    contours, _ = cv2.findContours(edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    
    # Sort contours by area (largest to smallest)
    contours = sorted(contours, key=cv2.contourArea, reverse=True)

    # We need to identify the table structure (rows and columns)
    table_structure = []

    for contour in contours:
        if cv2.contourArea(contour) > 500:  # Filter out small contours
            x, y, w, h = cv2.boundingRect(contour)
            table_crop = img[y:y+h, x:x+w]
            cell_text = pytesseract.image_to_string(table_crop, config='--psm 6')

            if cell_text.strip():  # Only consider non-empty cells
                table_structure.append((cell_text.strip(), (x, y, w, h)))  # Store cell text and its position
    
    return table_structure

def create_docx(text, tables):
    """ Create a .docx file with extracted text and tables. """
    doc = Document()

    # Add OCR text to document
    doc.add_paragraph("Extracted Text:")
    doc.add_paragraph(text)

    # If tables are detected, add them to the document
    if tables:
        doc.add_paragraph("\nDetected Tables:")

        # Sort cells by their Y and X position to create rows and columns
        sorted_cells = sorted(tables, key=lambda x: (x[1][1], x[1][0]))  # Sort by Y, then X

        # Create a list to store rows of cells
        rows = []
        current_row = []
        previous_y = sorted_cells[0][1][1]
        threshold = 20  # Distance threshold to group cells into rows
        
        for cell_text, (x, y, w, h) in sorted_cells:
            if abs(y - previous_y) < threshold:  # Same row if Y difference is small
                current_row.append(cell_text)
            else:
                rows.append(current_row)
                current_row = [cell_text]
            previous_y = y
        
        if current_row:  # Add last row if present
            rows.append(current_row)

        # Create a table in the Word document with the same structure
        num_columns = max(len(row) for row in rows)
        table = doc.add_table(rows=len(rows), cols=num_columns)

        # Populate the table with extracted text
        for row_idx, row in enumerate(rows):
            for col_idx, cell_text in enumerate(row):
                table.cell(row_idx, col_idx).text = cell_text

    # Save the docx file
    doc_path = 'static/output_with_table.docx'
    doc.save(doc_path)

    return doc_path

if __name__ == '__main__':
    app.run(debug=True)
